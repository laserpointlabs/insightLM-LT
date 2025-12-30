#!/usr/bin/env python3
"""
Workbook RAG Server - On-Demand Reading with Content Search
Searches document content (PDFs, Word, text) using smart text matching.
"""
import os
import sys
import json
import re
from pathlib import Path
from typing import List, Dict, Any, Optional

# Get data directory from environment
DATA_DIR = os.environ.get("INSIGHTLM_DATA_DIR", "")

# In-memory cache for documents (keyed by filepath, stores (content, mtime))
_document_cache: Dict[str, tuple[str, float]] = {}
_cache_timestamp: Optional[float] = None

def _prune_document_cache() -> None:
    """Remove cache entries for files that no longer exist."""
    global _document_cache
    try:
        dead_keys = []
        for k in list(_document_cache.keys()):
            try:
                if not Path(k).exists():
                    dead_keys.append(k)
            except Exception:
                dead_keys.append(k)
        for k in dead_keys:
            _document_cache.pop(k, None)
    except Exception:
        # Cache is best-effort only
        pass

def _compute_cache_stamp(workbooks_dir: Path) -> float:
    """Compute a stamp that changes when workbook metadata changes.

    We cannot rely on the parent directory mtime alone on all platforms/filesystems.
    """
    stamp = 0.0
    try:
        stamp = max(stamp, workbooks_dir.stat().st_mtime)
    except Exception:
        pass
    try:
        for meta in workbooks_dir.glob("*/workbook.json"):
            try:
                stamp = max(stamp, meta.stat().st_mtime)
            except Exception:
                continue
    except Exception:
        pass
    return stamp

def clear_cache(workbook_id: Optional[str] = None, file_path: Optional[str] = None) -> Dict[str, Any]:
    """Clear cache entries.

    Args:
        workbook_id: If set, clears entries under this workbook folder only.
        file_path: If set, clears a specific absolute file path (best-effort).
    """
    global _document_cache, _cache_timestamp
    removed = 0

    if file_path:
        k = str(Path(file_path))
        if k in _document_cache:
            _document_cache.pop(k, None)
            removed += 1
        return {"cleared": removed, "scope": "file"}

    if workbook_id:
        try:
            data_dir = Path(get_data_dir())
            wb_dir = data_dir / "workbooks" / workbook_id
            prefix = str(wb_dir)
            for k in list(_document_cache.keys()):
                if str(k).startswith(prefix):
                    _document_cache.pop(k, None)
                    removed += 1
        except Exception:
            # fallback: clear all
            removed = len(_document_cache)
            _document_cache.clear()
        return {"cleared": removed, "scope": "workbook"}

    removed = len(_document_cache)
    _document_cache.clear()
    _cache_timestamp = None
    return {"cleared": removed, "scope": "all"}

def get_data_dir() -> str:
    """Get application data directory"""
    if DATA_DIR:
        print(f"DEBUG: Using INSIGHTLM_DATA_DIR: {DATA_DIR}", file=sys.stderr, flush=True)
        return DATA_DIR

    if sys.platform == "win32":
        appdata = os.environ.get("APPDATA", "")
        if appdata:
            return str(Path(appdata) / "insightLM-LT")
    else:
        home = os.environ.get("HOME", "")
        if home:
            return str(Path(home) / ".config" / "insightLM-LT")

    raise ValueError("Could not determine data directory")


def extract_text_from_pdf(file_path: Path) -> str:
    """Extract text from PDF"""
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]
        reader = PdfReader(str(file_path))
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return '\n\n'.join(text_parts) if text_parts else ""
    except Exception as e:
        return f"Error extracting PDF: {e}"


def extract_text_from_docx(file_path: Path) -> str:
    """Extract text from DOCX"""
    try:
        from docx import Document  # type: ignore[import-not-found]
        doc = Document(str(file_path))
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(' | '.join(row_text))
        return '\n\n'.join(text_parts) if text_parts else ""
    except Exception as e:
        return f"Error extracting DOCX: {e}"


def extract_text_from_excel(file_path: Path) -> str:
    """Extract text from Excel (XLSX, XLS)"""
    try:
        import pandas as pd

        # Read all sheets
        excel_file = pd.ExcelFile(str(file_path))
        text_parts = []

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)

            # Add sheet name as header
            text_parts.append(f"=== Sheet: {sheet_name} ===")

            # Convert dataframe to text representation
            # Include column headers and all rows
            text_parts.append(df.to_string(index=False))
            text_parts.append("")  # Empty line between sheets

        return '\n\n'.join(text_parts) if text_parts else ""
    except Exception as e:
        return f"Error extracting Excel: {e}"


def extract_text_from_insight_sheet(file_path: Path) -> str:
    """Extract text from Insight Sheet (.is) file with formulas visible"""
    try:
        content = file_path.read_text(encoding='utf-8')
        data = json.loads(content)
        
        text_parts = []
        text_parts.append(f"Spreadsheet: {file_path.name}")
        
        if 'metadata' in data:
            metadata = data['metadata']
            if 'name' in metadata:
                text_parts.append(f"Name: {metadata['name']}")
            if 'workbook_id' in metadata:
                text_parts.append(f"Workbook ID: {metadata['workbook_id']}")
        
        text_parts.append("")  # Empty line
        
        # Process each sheet
        sheets = data.get('sheets', [])
        for sheet in sheets:
            sheet_name = sheet.get('name', 'Sheet1')
            sheet_id = sheet.get('id', 'sheet1')
            cells = sheet.get('cells', {})
            formats = sheet.get('formats', {})
            
            text_parts.append(f"=== Sheet: {sheet_name} ===")
            
            if not cells:
                text_parts.append("(Empty sheet)")
                text_parts.append("")
                continue
            
            # Extract all cell data with formulas visible
            cell_data = []
            for cell_ref, cell_info in cells.items():
                if isinstance(cell_info, dict):
                    # Handle nested value structure from Luckysheet format
                    # Format can be either:
                    # 1. { "value": 123, "formula": "=A1*2" } (simple format)
                    # 2. { "value": { "v": 123, "f": "=A1*2", "m": "123" } } (Luckysheet format)
                    
                    if 'value' in cell_info and isinstance(cell_info['value'], dict):
                        # Luckysheet format: nested value object
                        value_obj = cell_info['value']
                        formula = value_obj.get('f', '')  # Formula in 'f' field
                        value = value_obj.get('v', value_obj.get('m', ''))  # Value in 'v' or 'm' field
                        
                        if formula:
                            # Formula cell - show formula AND calculated value
                            cell_data.append(f"Cell {cell_ref}: {formula} (formula, calculated value: {value})")
                        else:
                            # Value cell
                            cell_data.append(f"Cell {cell_ref}: {value}")
                    else:
                        # Simple format: { "value": 123, "formula": "=A1*2" }
                        value = cell_info.get('value', '')
                        formula = cell_info.get('formula', '')
                        
                        if formula:
                            # Formula cell - show formula AND calculated value
                            cell_data.append(f"Cell {cell_ref}: {formula} (formula, calculated value: {value})")
                        else:
                            # Value cell
                            cell_data.append(f"Cell {cell_ref}: {value}")
                else:
                    # Simple value (direct value, not a dict)
                    cell_data.append(f"Cell {cell_ref}: {cell_info}")
            
            # Sort cells by reference (A1, A2, B1, etc.)
            def cell_sort_key(ref: str):
                # Extract column (letters) and row (numbers)
                match = re.match(r'([A-Z]+)(\d+)', ref)
                if match:
                    col = match.group(1)
                    row = int(match.group(2))
                    # Convert column to number (A=1, B=2, ..., Z=26, AA=27, etc.)
                    col_num = sum((ord(c) - ord('A') + 1) * (26 ** i) for i, c in enumerate(reversed(col)))
                    return (row, col_num)
                return (0, 0)
            
            def get_cell_ref_from_line(line: str) -> str:
                match = re.search(r'Cell ([A-Z]+\d+)', line)
                return match.group(1) if match else 'A1'
            
            cell_data.sort(key=lambda x: cell_sort_key(get_cell_ref_from_line(x)))
            
            text_parts.extend(cell_data)
            text_parts.append("")  # Empty line between sheets
        
        # Add formula dependencies summary
        formula_cells = []
        for sheet in sheets:
            cells = sheet.get('cells', {})
            for cell_ref, cell_info in cells.items():
                if isinstance(cell_info, dict):
                    # Extract formula from either format
                    formula = None
                    if 'value' in cell_info and isinstance(cell_info['value'], dict):
                        # Luckysheet format: formula in value.f
                        formula = cell_info['value'].get('f', '')
                    else:
                        # Simple format: formula in cell_info.formula
                        formula = cell_info.get('formula', '')
                    
                    if formula:
                        dependencies = extract_cell_references(formula)
                        formula_cells.append(f"{cell_ref}: {formula} (depends on: {', '.join(dependencies)})")
        
        if formula_cells:
            text_parts.append("=== Formulas ===")
            text_parts.extend(formula_cells)
        
        # Extract conditional formatting rules
        conditional_format_rules = []
        conditional_formats = sheet.get('conditionalFormats', {})
        if conditional_formats:
            for rule_key, rule in conditional_formats.items():
                if not isinstance(rule, dict):
                    continue
                    
                # Rule structure varies, but common fields:
                # - type: 'cellValue', 'formula', 'textContains', etc.
                # - cellrange: cell range like "A1:A10" or [{r:0, c:0}, {r:0, c:0}]
                # - condition: operator like '>', '<', '=', 'between', etc.
                # - value: threshold value(s)
                # - format: formatting object with bg (background), fc (font color), etc.
                
                rule_type = rule.get('type', rule.get('conditionType', 'unknown'))
                cell_range = rule.get('cellrange', rule.get('range', rule_key))
                
                # Handle cell range format (could be string "A1:A10" or array)
                if isinstance(cell_range, list) and len(cell_range) >= 2:
                    # Convert array format to string
                    start = cell_range[0]
                    end = cell_range[1] if len(cell_range) > 1 else start
                    if isinstance(start, dict) and 'r' in start and 'c' in start:
                        # Convert row/col to cell reference
                        start_col = chr(65 + start['c']) if start['c'] < 26 else 'A'  # Simple conversion
                        start_row = start['r'] + 1
                        end_col = chr(65 + end['c']) if end['c'] < 26 else 'A'
                        end_row = end['r'] + 1
                        cell_range = f"{start_col}{start_row}:{end_col}{end_row}"
                
                condition = rule.get('condition', rule.get('operator', ''))
                value = rule.get('value', rule.get('threshold', ''))
                format_info = rule.get('format', rule.get('style', {}))
                
                # Format the rule description based on type
                if rule_type == 'cellValue' or rule_type == 'number':
                    # Example: "Cell value > 100" -> red background
                    if condition and value:
                        condition_desc = f"value {condition} {value}"
                    elif condition:
                        condition_desc = f"value {condition}"
                    else:
                        condition_desc = f"value check"
                elif rule_type == 'formula':
                    # Formula-based condition
                    formula = rule.get('formula', condition)
                    condition_desc = f"formula: {formula}"
                elif rule_type == 'textContains':
                    condition_desc = f"text contains '{value}'" if value else "text contains"
                elif rule_type == 'duplicate':
                    condition_desc = "duplicate values"
                elif rule_type == 'unique':
                    condition_desc = "unique values"
                else:
                    condition_desc = condition if condition else f"{rule_type} condition"
                
                # Extract formatting details (color, style, etc.)
                format_desc = []
                if isinstance(format_info, dict):
                    # Background color
                    bg_color = format_info.get('bg', format_info.get('backgroundColor', format_info.get('backColor')))
                    if bg_color:
                        # Handle hex colors (#FF0000) or color names
                        if isinstance(bg_color, str):
                            format_desc.append(f"background: {bg_color}")
                        elif isinstance(bg_color, dict) and 'rgb' in bg_color:
                            format_desc.append(f"background: rgb({bg_color['rgb']})")
                    
                    # Font color
                    font_color = format_info.get('fc', format_info.get('fontColor', format_info.get('foreColor')))
                    if font_color:
                        if isinstance(font_color, str):
                            format_desc.append(f"font color: {font_color}")
                        elif isinstance(font_color, dict) and 'rgb' in font_color:
                            format_desc.append(f"font color: rgb({font_color['rgb']})")
                    
                    # Text style
                    if format_info.get('bl', format_info.get('bold', False)):
                        format_desc.append("bold")
                    if format_info.get('it', format_info.get('italic', False)):
                        format_desc.append("italic")
                    if format_info.get('un', format_info.get('underline', False)):
                        format_desc.append("underline")
                
                format_str = ", ".join(format_desc) if format_desc else "formatting applied"
                conditional_format_rules.append(
                    f"Conditional format on {cell_range}: when {condition_desc} -> {format_str}"
                )
        
        if conditional_format_rules:
            text_parts.append("=== Conditional Formatting ===")
            text_parts.extend(conditional_format_rules)
        
        return '\n'.join(text_parts)
    except json.JSONDecodeError as e:
        return f"Error parsing Insight Sheet JSON: {e}"
    except Exception as e:
        return f"Error extracting Insight Sheet: {e}"


def extract_cell_references(formula: str) -> List[str]:
    """Extract cell references from a formula (e.g., A1, B2, etc.)"""
    if not formula.startswith("="):
        return []
    
    # Pattern to match cell references like A1, B2, AA10, etc.
    pattern = r'\b([A-Z]+)(\d+)\b'
    matches = re.findall(pattern, formula)
    return [f"{col}{row}" for col, row in matches]


def read_file(file_path: Path) -> str:
    """Read any file type"""
    ext = file_path.suffix.lower()

    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif ext in ['.xlsx', '.xls']:
        return extract_text_from_excel(file_path)
    elif ext == '.is':
        # Insight Sheet format - extract with formulas visible
        return extract_text_from_insight_sheet(file_path)
    else:
        # Text file (includes .csv, .txt, .md, etc.)
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                return file_path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        return f"Could not decode file: {file_path}"


def _iter_grep_matches(
    content: str,
    pattern: str,
    *,
    regex: bool,
    case_sensitive: bool,
    max_matches: int,
) -> List[Dict[str, Any]]:
    """Find matches in content with best-effort line/col and a small snippet.

    Notes:
    - Works for any extracted text, not only "true" text files.
    - Results are deterministic (scan left-to-right).
    """
    if not pattern:
        return []

    matches: List[Dict[str, Any]] = []
    content_str = content if isinstance(content, str) else str(content)

    if regex:
        flags = re.MULTILINE
        if not case_sensitive:
            flags |= re.IGNORECASE

        try:
            compiled = re.compile(pattern, flags)
        except re.error as e:
            raise ValueError(f"Invalid regex: {e}")

        for m in compiled.finditer(content_str):
            if len(matches) >= max_matches:
                break
            start = m.start()
            end = m.end()
            matches.append({"start": start, "end": end})
    else:
        haystack = content_str if case_sensitive else content_str.lower()
        needle = pattern if case_sensitive else pattern.lower()
        i = 0
        nlen = len(needle)
        if nlen == 0:
            return []

        while True:
            if len(matches) >= max_matches:
                break
            idx = haystack.find(needle, i)
            if idx == -1:
                break
            matches.append({"start": idx, "end": idx + nlen})
            i = idx + (1 if nlen == 0 else nlen)

    # Enrich with line/col + snippet (best-effort)
    if not matches:
        return matches

    # Precompute line starts for O(log n) mapping via linear scan (content sizes are small enough)
    # We keep it simple and deterministic.
    line_starts = [0]
    for idx, ch in enumerate(content_str):
        if ch == "\n":
            line_starts.append(idx + 1)

    def _line_col(pos: int) -> tuple[int, int]:
        # Find rightmost line_start <= pos (linear scan acceptable; line_starts size bounded by content)
        line = 0
        for j, s in enumerate(line_starts):
            if s <= pos:
                line = j
            else:
                break
        col = pos - line_starts[line]
        return (line + 1, col + 1)  # 1-based

    SNIPPET_RADIUS = 80
    enriched: List[Dict[str, Any]] = []
    for m in matches:
        start = int(m["start"])
        end = int(m["end"])
        line, col = _line_col(start)
        snippet_start = max(0, start - SNIPPET_RADIUS)
        snippet_end = min(len(content_str), end + SNIPPET_RADIUS)
        snippet = content_str[snippet_start:snippet_end]
        if snippet_start > 0:
            snippet = "..." + snippet
        if snippet_end < len(content_str):
            snippet = snippet + "..."
        enriched.append(
            {
                "start": start,
                "end": end,
                "line": line,
                "col": col,
                "snippet": snippet,
            }
        )

    return enriched


def grep_workbooks(
    pattern: str,
    *,
    regex: bool = False,
    case_sensitive: bool = False,
    max_results: int = 50,
    max_matches_per_file: int = 20,
    workbook_ids: Optional[List[str]] = None,
) -> Dict[str, Any]:
    """Grep-like search across workbook documents (workbooks only).

    Returns structured results suitable for LLM tool usage.
    """
    if max_results <= 0:
        return {"pattern": pattern, "regex": regex, "case_sensitive": case_sensitive, "results": [], "truncated": False}

    # Defensive caps
    MAX_PATTERN_LEN = 500
    if pattern is None:
        pattern = ""
    if len(pattern) > MAX_PATTERN_LEN:
        raise ValueError(f"Pattern too long (max {MAX_PATTERN_LEN} chars)")

    if max_matches_per_file <= 0:
        max_matches_per_file = 1

    docs = get_all_workbook_documents(workbook_ids)
    results: List[Dict[str, Any]] = []
    truncated = False

    # Deterministic ordering: workbook_id, then path
    docs.sort(key=lambda d: (d.get("workbook_id", ""), d.get("path", ""), d.get("filename", "")))

    for doc in docs:
        if len(results) >= max_results:
            truncated = True
            break

        content = doc.get("content", "")
        if not content:
            continue

        matches = _iter_grep_matches(
            content,
            pattern,
            regex=bool(regex),
            case_sensitive=bool(case_sensitive),
            max_matches=int(max_matches_per_file),
        )
        if not matches:
            continue

        results.append(
            {
                "workbook_id": doc.get("workbook_id", ""),
                "workbook_name": doc.get("workbook_name", ""),
                "filename": doc.get("filename", ""),
                "path": doc.get("path", ""),
                "full_path": doc.get("filepath", ""),
                "match_count": len(matches),
                "matches": matches,
            }
        )

    return {
        "pattern": pattern,
        "regex": bool(regex),
        "case_sensitive": bool(case_sensitive),
        "max_results": int(max_results),
        "max_matches_per_file": int(max_matches_per_file),
        "workbook_ids": workbook_ids,
        "results": results,
        "truncated": truncated,
    }


def get_all_workbook_documents(workbook_ids: Optional[List[str]] = None) -> List[Dict[str, Any]]:
    """Scan workbooks and return document metadata with content (cached).

    If workbook_ids is provided, only those workbook directory names are scanned.
    """
    global _document_cache, _cache_timestamp

    workbook_id_allowlist = set(workbook_ids) if workbook_ids else None
    documents = []
    data_dir = Path(get_data_dir())
    workbooks_dir = data_dir / "workbooks"

    print(f"DEBUG: Looking for workbooks in: {workbooks_dir}", file=sys.stderr, flush=True)
    print(f"DEBUG: Workbooks dir exists: {workbooks_dir.exists()}", file=sys.stderr, flush=True)

    if not workbooks_dir.exists():
        print(f"DEBUG: Workbooks directory not found: {workbooks_dir}", file=sys.stderr, flush=True)
        return []

    # Prune deleted files from cache so we never return "legacy" content.
    _prune_document_cache()

    # Check if cache is still valid (re-scan if workbook metadata changed).
    try:
        current_stamp = _compute_cache_stamp(workbooks_dir)
        if _cache_timestamp is not None and current_stamp <= _cache_timestamp:
            # Cache is valid, but we still need to rebuild the documents list from cache
            pass
        else:
            # Invalidate cache if directory changed
            _document_cache.clear()
            _cache_timestamp = current_stamp
    except:
        _document_cache.clear()
        _cache_timestamp = None

    for workbook_dir in workbooks_dir.iterdir():
        if not workbook_dir.is_dir():
            continue
        if workbook_id_allowlist is not None and workbook_dir.name not in workbook_id_allowlist:
            continue

        metadata_path = workbook_dir / "workbook.json"
        if not metadata_path.exists():
            continue

        try:
            with open(metadata_path, 'r', encoding='utf-8') as f:
                metadata = json.load(f)
        except:
            continue

        workbook_name = metadata.get('name', workbook_dir.name)

        for doc in metadata.get('documents', []):
            filename = doc.get('filename', '')
            relative_path = doc.get('path', f"documents/{filename}")
            file_path = workbook_dir / relative_path

            if not file_path.exists():
                continue

            # Check cache first
            content = ""
            cache_key = str(file_path)
            try:
                file_mtime = file_path.stat().st_mtime

                if cache_key in _document_cache:
                    cached_content, cached_mtime = _document_cache[cache_key]
                    if file_mtime <= cached_mtime:
                        # Cache hit - use cached content
                        content = cached_content
                    else:
                        # File changed - re-extract
                        _document_cache.pop(cache_key, None)
            except:
                file_mtime = 0

            # Extract text if not in cache
            if not content:
                content = read_file(file_path)

                # Cache the extracted content
                if content and not content.startswith("Error"):
                    try:
                        _document_cache[cache_key] = (content, file_mtime)
                    except:
                        pass

            if content and not content.startswith("Error"):
                documents.append({
                    'workbook_id': workbook_dir.name,
                    'workbook_name': workbook_name,
                    'filename': filename,
                    'path': str(relative_path),
                    'filepath': str(file_path),
                    'content': content,
                })

    return documents


def search_workbooks(query: str, limit: int = 20) -> List[Dict[str, Any]]:
    """Search for files matching query - returns file metadata only (for backward compatibility)"""
    documents = get_all_workbook_documents()

    if not documents:
        return []

    query_lower = query.lower()
    matching_docs = []

    for doc in documents:
        # Calculate relevance score (simple text matching)
        score = 0
        content_lower = doc["content"].lower()

        # Exact phrase match in content (high score)
        if query_lower in content_lower:
            score += 20

        # Word matching - each word adds points
        query_words = query_lower.split()
        word_matches = sum(1 for word in query_words if len(word) > 2 and word in content_lower)
        score += word_matches * 3

        # Filename match (medium score)
        filename_matches = sum(1 for word in query_words if word in doc["filename"].lower())
        score += filename_matches * 5

        # Workbook name match (low score)
        workbook_matches = sum(1 for word in query_words if word in doc["workbook_name"].lower())
        score += workbook_matches * 2

        # Special boost for PDFs (they often have important docs)
        if doc["filename"].lower().endswith(".pdf") and score > 0:
            score += 2

        if score > 0:
            matching_docs.append((score, doc))

    # Sort by relevance score
    matching_docs.sort(key=lambda x: x[0], reverse=True)

    # Return top results (file metadata only for backward compatibility)
    results = []
    for score, doc in matching_docs[:limit]:
        results.append({
            'workbook_id': doc['workbook_id'],
            'workbook_name': doc['workbook_name'],
            'filename': doc['filename'],
            'path': doc['path'],
            'full_path': doc['filepath'],
            'match_type': 'content',
            'relevance_score': score
        })

    return results


def extract_context_chunks(content: str, key_terms: List[str], chunk_size: int = 1000, max_chunks: int = 5) -> List[tuple[int, str]]:
    """
    Extract context chunks around keyword matches in document content.
    Returns chunks of text surrounding where keywords are found.

    Args:
        content: Full document content
        key_terms: List of keywords to search for
        chunk_size: Size of chunk (chars before + after keyword)
        max_chunks: Maximum number of chunks to return

    Returns:
        List of (position, chunk_text) tuples
    """
    content_lower = content.lower()
    chunks = []
    positions_covered = set()

    # Find all keyword positions
    for term in key_terms:
        if not term or len(term) < 3:
            continue

        # Find all occurrences of this term
        pattern = r'\b' + re.escape(term) + r'\b'
        for match in re.finditer(pattern, content_lower):
            pos = match.start()

            # Skip if we already have a chunk covering this position
            if any(abs(pos - covered) < chunk_size // 2 for covered in positions_covered):
                continue

            # Extract chunk around this position
            start = max(0, pos - chunk_size // 2)
            end = min(len(content), pos + chunk_size // 2)

            # Try to break at word boundaries
            if start > 0:
                # Find previous space
                space_before = content.rfind(' ', start - 100, start)
                if space_before > 0:
                    start = space_before + 1

            if end < len(content):
                # Find next space
                space_after = content.find(' ', end, end + 100)
                if space_after > 0:
                    end = space_after

            chunk_text = content[start:end].strip()

            # Add ellipsis if not at document boundaries
            if start > 0:
                chunk_text = "..." + chunk_text
            if end < len(content):
                chunk_text = chunk_text + "..."

            chunks.append((pos, chunk_text))
            positions_covered.add(pos)

            if len(chunks) >= max_chunks:
                break

        if len(chunks) >= max_chunks:
            break

    # Sort chunks by position in document
    chunks.sort(key=lambda x: x[0])

    return chunks


def search_workbooks_with_content(query: str, limit: int = 5, workbook_ids: Optional[List[str]] = None) -> str:
    """Search workbook documents using text matching and return FULL content.
    Returns only files that match the query (score >= 3) plus up to 2 relevant sibling files per workbook.
    Limits to 5 files per workbook to avoid overwhelming results.
    """
    documents = get_all_workbook_documents(workbook_ids)

    if not documents:
        return "No workbook documents found."

    # Extract key terms from query (remove common words and punctuation)
    query_lower = query.lower()
    # Remove punctuation from query
    query_clean = re.sub(r'[^\w\s]', ' ', query_lower)
    common_words = {'who', 'are', 'the', 'in', 'a', 'an', 'and', 'or', 'but', 'is', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'should', 'could', 'may', 'might', 'must', 'can', 'this', 'that', 'these', 'those', 'what', 'when', 'where', 'why', 'how'}
    query_words = [w for w in query_clean.split() if w not in common_words and len(w) > 2]

    # If we have key terms, require at least one key term match for a file to be included
    # Also include word stems for better matching (e.g., "cabins" -> "cabin")
    key_terms = query_words if query_words else [query_lower]
    # Add stemmed versions (simple: remove trailing 's')
    stemmed_terms = []
    for term in key_terms:
        stemmed_terms.append(term)
        if term.endswith('s') and len(term) > 3:
            stemmed_terms.append(term[:-1])  # "cabins" -> "cabin"
    key_terms = list(set(stemmed_terms))  # Remove duplicates

    matching_docs = []

    for doc in documents:
        # Calculate relevance score
        score = 0
        content_lower = doc["content"].lower()
        filename_lower = doc["filename"].lower()

        # Check if document contains at least one key term (required for inclusion)
        # Also check for partial filename matches (e.g., "spreadsheet-2025-12-12" matches "spreadsheet-2025-12-12T19-46-01.is")
        has_key_term = False
        for term in key_terms:
            if term in content_lower or term in filename_lower:
                has_key_term = True
                break
        
        # Special handling: if query contains a filename pattern (e.g., "spreadsheet-2025-12-12T19-41-01"),
        # also match similar filenames (same date prefix)
        if not has_key_term:
            # Extract date prefix from query (e.g., "spreadsheet-2025-12-12" from "spreadsheet-2025-12-12T19-41-01")
            date_prefix_match = re.search(r'([a-z0-9_-]+-\d{4}-\d{2}-\d{2})', query_lower)
            if date_prefix_match:
                date_prefix = date_prefix_match.group(1)
                # Check if filename starts with this prefix
                if filename_lower.startswith(date_prefix):
                    has_key_term = True
                    score += 10  # Boost score for filename prefix match

        # Skip documents that don't contain any key terms
        if not has_key_term:
            continue

        # Exact phrase match in content (high score)
        if query_lower in content_lower:
            score += 20
        elif any(term in content_lower for term in key_terms):
            # Key term match (but not exact phrase)
            score += 15

        # Word matching - use key terms with word boundaries
        if key_terms:
            word_matches = sum(1 for word in key_terms if re.search(r'\b' + re.escape(word) + r'\b', content_lower))
            score += word_matches * 5  # Higher weight for key terms

        # Filename match (very important - boost significantly)
        filename_matches = sum(1 for term in key_terms if term in filename_lower)
        score += filename_matches * 10  # Strong filename match boost

        # Workbook name match
        workbook_matches = sum(1 for term in key_terms if term in doc["workbook_name"].lower())
        score += workbook_matches * 3

        # Special boost for PDFs
        if doc["filename"].lower().endswith(".pdf") and score > 0:
            score += 2

        # Only include files with meaningful matches (score >= 8 for stricter filtering)
        # This ensures we only get files that really match the query
        if score >= 8:
            matching_docs.append((score, doc))

    # Sort by relevance score
    matching_docs.sort(key=lambda x: x[0], reverse=True)

    if not matching_docs:
        available = "\n".join([f"- {d['filename']} ({d['workbook_name']})" for d in documents[:20]])
        return f"No matches found for '{query}'.\n\nAvailable documents:\n{available}"

    # Only return files with scores close to the top score
    # This prevents returning all NDAs when only one matches
    top_score = matching_docs[0][0]
    score_threshold = top_score * 0.9  # Must be within 90% of top score (stricter)

    filtered_docs = [(score, doc) for score, doc in matching_docs if score >= score_threshold]

    # If no files pass 90% threshold, return only the top file
    if len(filtered_docs) == 0:
        filtered_docs = [matching_docs[0]]

    # Limit total results
    MAX_TOTAL_FILES = 2  # Return at most 2 files total
    filtered_docs = filtered_docs[:MAX_TOTAL_FILES]

    # Format results with context-aware chunks
    results = []
    for score, doc in filtered_docs:
        # Extract context chunks around keywords instead of returning full content
        chunks = extract_context_chunks(doc["content"], key_terms, chunk_size=1000, max_chunks=5)

        if len(chunks) == 0:
            # No specific keyword positions found, return beginning of file
            max_chars = 2000
            content_preview = doc["content"][:max_chars]
            if len(doc["content"]) > max_chars:
                content_preview += f"\n\n[Preview only - document is {len(doc['content'])} characters total. Use read_workbook_file to get the complete document.]"

            results.append(f"""**{doc['filename']}** ({doc['workbook_name']})
Workbook ID: {doc['workbook_id']}
Path: {doc['path']}
Relevance Score: {score}
Total Length: {len(doc['content'])} characters

=== PREVIEW (first 2,000 chars) ===
{content_preview}

---
""")
        else:
            # Return context chunks around keyword matches
            chunks_text = "\n\n".join([f"[Excerpt {i+1} - Position {pos:,}]\n{chunk}"
                                        for i, (pos, chunk) in enumerate(chunks)])

            total_chunk_chars = sum(len(chunk) for _, chunk in chunks)

            results.append(f"""**{doc['filename']}** ({doc['workbook_name']})
Workbook ID: {doc['workbook_id']}
Path: {doc['path']}
Relevance Score: {score}
Total Length: {len(doc['content'])} characters
Chunks: {len(chunks)} excerpts ({total_chunk_chars} characters total)

=== CONTEXT EXCERPTS (around keywords: {', '.join(key_terms[:5])}) ===
{chunks_text}

[NOTE: This shows only relevant excerpts. To read the complete document, use read_workbook_file with the workbook ID and path above.]

---
""")

    return "\n".join(results)


def read_workbook_file(workbook_id: str, file_path: str) -> str:
    """Read a specific file from a workbook"""
    data_dir = Path(get_data_dir())
    full_path = data_dir / "workbooks" / workbook_id / file_path

    if not full_path.exists():
        return f"File not found: {file_path}"

    return read_file(full_path)


def list_all_files(workbook_ids: Optional[List[str]] = None) -> List[Dict[str, Any]]:
    """List all files in workbooks.

    If workbook_ids is provided, only those workbook directory names are included.
    """
    files = []
    data_dir = Path(get_data_dir())
    workbooks_dir = data_dir / "workbooks"

    if not workbooks_dir.exists():
        return []

    workbook_id_allowlist = set(workbook_ids) if workbook_ids else None

    for workbook_dir in workbooks_dir.iterdir():
        if not workbook_dir.is_dir():
            continue
        if workbook_id_allowlist is not None and workbook_dir.name not in workbook_id_allowlist:
            continue

        metadata_path = workbook_dir / "workbook.json"
        if not metadata_path.exists():
            continue

        try:
            with open(metadata_path, 'r', encoding='utf-8') as f:
                metadata = json.load(f)
        except:
            continue

        workbook_name = metadata.get('name', workbook_dir.name)

        for doc in metadata.get('documents', []):
            files.append({
                'workbook_id': workbook_dir.name,
                'workbook_name': workbook_name,
                'filename': doc.get('filename', ''),
                'path': doc.get('path', ''),
                'added_at': doc.get('addedAt', '')
            })

    return files


def handle_request(request: dict) -> dict:
    """Handle MCP requests"""
    method = request.get('method', '')
    params = request.get('params', {})

    try:
        # MCP protocol initialization
        if method == 'initialize':
            return {
                'jsonrpc': '2.0',
                'id': request.get('id'),
                'result': {
                    'protocolVersion': '2024-11-05',
                    'capabilities': {
                        'tools': {
                            'listChanged': True
                        }
                    },
                    'serverInfo': {
                        'name': 'workbook-rag',
                        'version': '1.0.0'
                    }
                }
            }
        
        # MCP tools/list - expose available tools
        elif method == 'tools/list':
            return {
                'jsonrpc': '2.0',
                'id': request.get('id'),
                'result': {
                    'tools': [
                        {
                            'name': 'rag_search_content',
                            'description': 'Search document CONTENT (not just filenames) across all workbooks. This searches inside PDFs, Word docs, spreadsheets, and text files. Returns full content of matching files with context chunks.',
                            'inputSchema': {
                                'type': 'object',
                                'properties': {
                                    'query': {
                                        'type': 'string',
                                        'description': 'What to search for (e.g., "BSEO", "authentication", "expiration dates", "NDA terms", or filename patterns)'
                                    },
                                    'limit': {
                                        'type': 'number',
                                        'description': 'Maximum number of files to return (default: 5)'
                                    },
                                    'workbook_ids': {
                                        'type': 'array',
                                        'items': { 'type': 'string' },
                                        'description': 'Optional: limit search to these workbook IDs (directory names)'
                                    }
                                },
                                'required': ['query']
                            }
                        },
                        {
                            'name': 'rag_list_files',
                            'description': 'List all files in all workbooks with their metadata (workbook ID, workbook name, filename, path)',
                            'inputSchema': {
                                'type': 'object',
                                'properties': {
                                    'workbook_ids': {
                                        'type': 'array',
                                        'items': { 'type': 'string' },
                                        'description': 'Optional: limit listing to these workbook IDs (directory names)'
                                    }
                                },
                                'required': []
                            }
                        },
                        {
                            'name': 'rag_read_file',
                            'description': 'Read the full contents of a specific file from a workbook',
                            'inputSchema': {
                                'type': 'object',
                                'properties': {
                                    'workbook_id': {
                                        'type': 'string',
                                        'description': 'The ID of the workbook (directory name, e.g., "ac1000-main-project")'
                                    },
                                    'file_path': {
                                        'type': 'string',
                                        'description': 'The relative path to the file within the workbook (e.g., "documents/file.txt")'
                                    }
                                },
                                'required': ['workbook_id', 'file_path']
                            }
                        }
                        ,
                        {
                            'name': 'rag_clear_cache',
                            'description': 'Clear the server-side document content cache (useful after deletes/moves to ensure no stale results)',
                            'inputSchema': {
                                'type': 'object',
                                'properties': {
                                    'workbook_id': {
                                        'type': 'string',
                                        'description': 'Optional: clear cache entries only for this workbook id'
                                    },
                                    'file_path': {
                                        'type': 'string',
                                        'description': 'Optional: clear cache entry for this absolute file path'
                                    }
                                },
                                'required': []
                            }
                        }
                        ,
                        {
                            'name': 'rag_grep',
                            'description': 'Grep-like pattern search across workbook documents (workbooks only). Returns structured match locations + snippets. Use regex=false for literal strings; regex=true for regular expressions.',
                            'inputSchema': {
                                'type': 'object',
                                'properties': {
                                    'pattern': {
                                        'type': 'string',
                                        'description': 'Pattern to search for. Interpreted literally by default unless regex=true.'
                                    },
                                    'regex': {
                                        'type': 'boolean',
                                        'description': 'If true, treat pattern as a regular expression. Default: false.'
                                    },
                                    'case_sensitive': {
                                        'type': 'boolean',
                                        'description': 'If true, match case-sensitively. Default: false.'
                                    },
                                    'max_results': {
                                        'type': 'number',
                                        'description': 'Maximum number of files to return (default: 50).'
                                    },
                                    'max_matches_per_file': {
                                        'type': 'number',
                                        'description': 'Maximum matches to return per file (default: 20).'
                                    },
                                    'workbook_ids': {
                                        'type': 'array',
                                        'items': { 'type': 'string' },
                                        'description': 'Optional: limit grep to these workbook IDs (directory names).'
                                    }
                                },
                                'required': ['pattern']
                            }
                        }
                    ]
                }
            }
        
        # MCP tools/call - handle tool execution
        elif method == 'tools/call':
            tool_name = params.get('name', '')
            tool_args = params.get('arguments', {})
            
            if tool_name == 'rag_search_content':
                query = tool_args.get('query', '')
                limit = tool_args.get('limit', 5)
                workbook_ids = tool_args.get('workbook_ids', None)
                results = search_workbooks_with_content(query, limit, workbook_ids)
                return {
                    'jsonrpc': '2.0',
                    'id': request.get('id'),
                    'result': {'content': results}
                }
            
            elif tool_name == 'rag_list_files':
                workbook_ids = tool_args.get('workbook_ids', None)
                files = list_all_files(workbook_ids)
                return {
                    'jsonrpc': '2.0',
                    'id': request.get('id'),
                    'result': files
                }
            
            elif tool_name == 'rag_read_file':
                workbook_id = tool_args.get('workbook_id', '')
                file_path = tool_args.get('file_path', '')
                content = read_workbook_file(workbook_id, file_path)
                return {
                    'jsonrpc': '2.0',
                    'id': request.get('id'),
                    'result': {'content': content}
                }

            elif tool_name == 'rag_clear_cache':
                workbook_id = tool_args.get('workbook_id', None)
                file_path = tool_args.get('file_path', None)
                result = clear_cache(workbook_id=workbook_id, file_path=file_path)
                return {
                    'jsonrpc': '2.0',
                    'id': request.get('id'),
                    'result': result
                }

            elif tool_name == 'rag_grep':
                pattern = tool_args.get('pattern', '')
                regex = tool_args.get('regex', False)
                case_sensitive = tool_args.get('case_sensitive', False)
                max_results = tool_args.get('max_results', 50)
                max_matches_per_file = tool_args.get('max_matches_per_file', 20)
                workbook_ids = tool_args.get('workbook_ids', None)
                result = grep_workbooks(
                    pattern,
                    regex=bool(regex),
                    case_sensitive=bool(case_sensitive),
                    max_results=int(max_results),
                    max_matches_per_file=int(max_matches_per_file),
                    workbook_ids=workbook_ids,
                )
                return {
                    'jsonrpc': '2.0',
                    'id': request.get('id'),
                    'result': result
                }
            
            else:
                return {
                    'jsonrpc': '2.0',
                    'id': request.get('id'),
                    'error': {
                        'code': -32601,
                        'message': f'Unknown tool: {tool_name}'
                    }
                }
        
        # Backward compatible: legacy rag/* methods
        elif method == 'rag/search':
            # Backward compatible: returns file metadata only
            query = params.get('query', '')
            limit = params.get('limit', 20)
            results = search_workbooks(query, limit)
            return {
                'jsonrpc': '2.0',
                'id': request.get('id'),
                'result': results
            }

        elif method == 'rag/search_content':
            # New: returns full content of matching files
            query = params.get('query', '')
            limit = params.get('limit', 5)
            workbook_ids = params.get('workbook_ids', None)
            results = search_workbooks_with_content(query, limit, workbook_ids)
            return {
                'jsonrpc': '2.0',
                'id': request.get('id'),
                'result': {'content': results}
            }

        elif method == 'rag/read_file':
            workbook_id = params.get('workbook_id', '')
            file_path = params.get('file_path', '')
            content = read_workbook_file(workbook_id, file_path)
            return {
                'jsonrpc': '2.0',
                'id': request.get('id'),
                'result': {'content': content}
            }

        elif method == 'rag/list_files':
            workbook_ids = params.get('workbook_ids', None)
            files = list_all_files(workbook_ids)
            return {
                'jsonrpc': '2.0',
                'id': request.get('id'),
                'result': files
            }

        elif method == 'rag/health':
            return {
                'jsonrpc': '2.0',
                'id': request.get('id'),
                'result': {'status': 'healthy', 'mode': 'on-demand-with-content-search'}
            }

        else:
            return {
                'jsonrpc': '2.0',
                'id': request.get('id'),
                'error': {'code': -32601, 'message': f'Unknown method: {method}'}
            }

    except Exception as e:
        return {
            'jsonrpc': '2.0',
            'id': request.get('id'),
            'error': {'code': -32603, 'message': str(e)}
        }


if __name__ == '__main__':
    # Send initialization message on startup (like workbook-dashboard)
    init_response = {
        "jsonrpc": "2.0",
        "id": 1,
        "result": {
            "protocolVersion": "2024-11-05",
            "serverInfo": {
                "name": "workbook-rag",
                "version": "1.0.0"
            },
            "capabilities": {
                "tools": {
                    "listChanged": True
                }
            }
        }
    }
    print(json.dumps(init_response), flush=True)
    
    # Handle requests
    for line in sys.stdin:
        try:
            request = json.loads(line.strip())
            response = handle_request(request)
            if response:
                print(json.dumps(response))
                sys.stdout.flush()
        except Exception as e:
            error_response = {
                'jsonrpc': '2.0',
                'id': request.get('id') if 'request' in locals() else None,
                'error': {
                    'code': -32603,
                    'message': str(e)
                }
            }
            print(json.dumps(error_response))
            sys.stdout.flush()
